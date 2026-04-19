"""
DOCX fallback pipeline for complex PDFs.

Pipeline:
  1. PDF → DOCX via pdf2docx
  2. Edit DOCX text (preserving formatting at run level)
  3. DOCX → PDF via docx2pdf or LibreOffice
"""

from __future__ import annotations

from pathlib import Path
from loguru import logger


def pdf_to_docx(pdf_path: str | Path, docx_path: str | Path) -> Path:
    """Convert PDF to DOCX using pdf2docx."""
    from pdf2docx import Converter

    pdf_path = Path(pdf_path)
    docx_path = Path(docx_path)
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info(f"Converting PDF → DOCX: {pdf_path} → {docx_path}")

    cv = Converter(str(pdf_path))
    cv.convert(str(docx_path))
    cv.close()

    logger.info("PDF → DOCX conversion complete")
    return docx_path


def edit_docx_text(
    docx_path: str | Path,
    replacements: dict[str, str],
    output_path: str | Path | None = None,
) -> Path:
    """
    Edit text in a DOCX file while preserving formatting.

    Operates at the *run* level to maintain font/size/color.
    """
    from docx import Document

    docx_path = Path(docx_path)
    output_path = Path(output_path) if output_path else docx_path

    logger.info(f"Editing DOCX: {docx_path} ({len(replacements)} replacements)")

    doc = Document(str(docx_path))

    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                # Replace at the run level to preserve formatting
                _replace_in_paragraph(paragraph, old_text, new_text)

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            _replace_in_paragraph(paragraph, old_text, new_text)

    doc.save(str(output_path))
    logger.info(f"DOCX saved: {output_path}")
    return output_path


def _replace_in_paragraph(paragraph, old_text: str, new_text: str):
    """Replace text across runs while preserving run-level formatting."""
    # Build full paragraph text and map char positions to runs
    full_text = ""
    char_to_run: list[tuple[int, int]] = []  # (run_idx, char_in_run)

    for run_idx, run in enumerate(paragraph.runs):
        for char_idx in range(len(run.text)):
            char_to_run.append((run_idx, char_idx))
        full_text += run.text

    # Find occurrence
    start = full_text.find(old_text)
    if start == -1:
        return

    end = start + len(old_text)

    # Determine which runs are affected
    if start >= len(char_to_run):
        return

    first_run_idx = char_to_run[start][0]
    last_run_idx = char_to_run[min(end - 1, len(char_to_run) - 1)][0]

    # Replace in first run, clear subsequent affected runs
    first_run = paragraph.runs[first_run_idx]
    first_char_in_run = char_to_run[start][1]

    # Text before the match in the first run
    prefix = first_run.text[:first_char_in_run]

    # Text after the match in the last run
    if end < len(char_to_run):
        last_run_idx_actual = char_to_run[end][0]
        last_char = char_to_run[end][1]
        suffix = paragraph.runs[last_run_idx_actual].text[last_char:]
    elif end == len(char_to_run):
        suffix = ""
        last_run_idx_actual = last_run_idx
    else:
        suffix = ""
        last_run_idx_actual = last_run_idx

    # Set first run's text
    first_run.text = prefix + new_text

    # Clear intermediate runs
    for r_idx in range(first_run_idx + 1, last_run_idx_actual + 1):
        if r_idx < len(paragraph.runs):
            paragraph.runs[r_idx].text = ""

    # If suffix was in a different run, append it
    if last_run_idx_actual != first_run_idx and suffix:
        first_run.text += suffix


def docx_to_pdf(docx_path: str | Path, pdf_path: str | Path) -> Path:
    """
    Convert DOCX back to PDF.

    Tries docx2pdf first, falls back to LibreOffice CLI.
    """
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info(f"Converting DOCX → PDF: {docx_path} → {pdf_path}")

    # Try docx2pdf (Windows/Mac with Word installed)
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        logger.info("DOCX → PDF via docx2pdf succeeded")
        return pdf_path
    except Exception as e:
        logger.warning(f"docx2pdf failed: {e}, trying LibreOffice...")

    # Fallback: LibreOffice CLI
    import subprocess

    try:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(pdf_path.parent),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode == 0:
            # LibreOffice outputs with same name but .pdf extension
            generated = pdf_path.parent / (docx_path.stem + ".pdf")
            if generated != pdf_path:
                generated.rename(pdf_path)
            logger.info("DOCX → PDF via LibreOffice succeeded")
            return pdf_path
        else:
            raise RuntimeError(f"LibreOffice error: {result.stderr}")
    except FileNotFoundError:
        raise RuntimeError(
            "Neither docx2pdf nor LibreOffice available for DOCX→PDF conversion"
        )


def fallback_pipeline(
    pdf_path: str | Path,
    replacements: dict[str, str],
    output_pdf_path: str | Path,
) -> Path:
    """
    Full DOCX fallback: PDF → DOCX → edit → PDF.
    """
    pdf_path = Path(pdf_path)
    output_pdf_path = Path(output_pdf_path)

    # Intermediate DOCX path
    docx_path = output_pdf_path.with_suffix(".docx")

    logger.info("Running DOCX fallback pipeline...")

    pdf_to_docx(pdf_path, docx_path)
    edit_docx_text(docx_path, replacements)
    docx_to_pdf(docx_path, output_pdf_path)

    # Clean up intermediate DOCX
    try:
        docx_path.unlink()
    except Exception:
        pass

    logger.info(f"DOCX fallback complete: {output_pdf_path}")
    return output_pdf_path
